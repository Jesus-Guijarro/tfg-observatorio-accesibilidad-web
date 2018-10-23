#AChecker

from selenium import webdriver

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

import io
import re

from docx import Document


document = Document()

document.add_heading('Document Title', 0)



#Se virtualiza una ventana de navegacion de Chrome

options = webdriver.ChromeOptions()

options.binary_location = '/usr/bin/google-chrome'
#options.add_argument('headless')

#Pruebas
options.add_argument('window-size=1200x600')

driver = webdriver.Chrome(chrome_options=options)

#Accedemos a la web de la herramienta de evaluacion
driver.get('https://achecker.ca/checker/index.php')
#driver.get("file:///home/jesus/Desktop/testACHECKER.html")


#Clase para crear condicion de espera a la hora exportar
class element_has_value(object):
 
  def __init__(self, locator, value):
    self.locator = locator
    self.value = value

  def __call__(self, driver):
    element = driver.find_element(*self.locator)   
    if self.value in element.get_attribute("value"):
        return element
    else:
        return False


wait = WebDriverWait(driver, 200)


elem =wait.until(EC.title_is(("IDI Web Accessibility Checker : Web Accessibility Checker")))

#Introducir URL
enlace = driver.find_element_by_css_selector('#checkuri')
enlace.clear()
enlace.send_keys('http://www.elmundo.es/internacional.html')

opciones=  driver.find_element_by_css_selector("#center-content > table > tbody > tr > td > div > form > div > fieldset > div:nth-child(6) > h2 > a")
opciones.click()

#Seleccionar opción WCAG 2.0
#Es necesario esperar ya que las opciones se generan con javascript
elem = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR,"#radio_gid_9"))) 
opcionWCAG2=  driver.find_element_by_css_selector("#radio_gid_9")
opcionWCAG2.click()

#Evaluar
boton= driver.find_element_by_css_selector("#validate_uri")
boton.click()


#Resultado
elem = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#report_file > div > label:nth-child(1)")))

numErrores=driver.find_element_by_css_selector("#AC_num_of_errors")
print(int(numErrores.text))

numProbables=driver.find_element_by_css_selector("#AC_num_of_likely")
print(int(numProbables.text))

numPotenciales=driver.find_element_by_css_selector("#AC_num_of_potential")
print(int(numPotenciales.text))


#Generar reporte
exportar = driver.find_element_by_css_selector("#validate_file_button")
exportar.click()

element = wait.until(element_has_value((By.CSS_SELECTOR, "#validate_file_button"), "Get File"))



#Comprobar si hay errores de X tipo o no
'''

if EC.presence_of_element_located((By.CSS_SELECTOR, "#AC_congrats_msg_for_errors > img")) == False:
    erroresConocidos=driver.find_element_by_id("AC_errors")
if EC.presence_of_element_located((By.CSS_SELECTOR, "#AC_congrats_msg_for_likely > img")) == False:
    erroresPosibles=driver.find_element_by_id("AC_likely_problems")
if EC.presence_of_element_located((By.CSS_SELECTOR, "#AC_congrats_msg_for_potential > img")) == False:

'''

if EC.presence_of_element_located((By.ID, "AC_errors")):
    erroresConocidos=driver.find_element_by_id("AC_errors")
if EC.presence_of_element_located((By.ID, "AC_likely_problems")):
    erroresPosibles=driver.find_element_by_id("AC_likely_problems")
if EC.presence_of_element_located((By.ID, "AC_potential_problems")):
    erroresPotenciales=driver.find_element_by_id("AC_potential_problems")

datosConocidos=str(erroresConocidos.get_attribute('textContent'))
datosPosibles=str(erroresPosibles.get_attribute('textContent'))
datosPotenciales=str(erroresPotenciales.get_attribute('textContent'))

#No son simples saltos de linea , son lineas con espacios en blanco diferentes, tabulados, etc
datosConocidos=datosConocidos.replace('  ','')
datosConocidos=datosConocidos.replace('\n','')
datosConocidos=datosConocidos.replace('        ','\n\n')
datosConocidos=datosConocidos.replace('Success Criteria','\n\nSuccess Criteria')
datosConocidos=datosConocidos.replace('Check','\n\n\tCheck')
datosConocidos=datosConocidos.replace('Repair','\n\n\tRepair')
datosConocidos=datosConocidos.replace('Line','\n\n\t\tLine')
datosConocidos=datosConocidos.replace('\t\t\t','')			


datosPotenciales=datosPotenciales.replace('  ','')
datosPotenciales=datosPotenciales.replace('\n','')
datosPotenciales=datosPotenciales.replace('        ','\n\n')
datosPotenciales=datosPotenciales.replace('Success Criteria','\n\nSuccess Criteria')
datosPotenciales=datosPotenciales.replace('Check','\n\n\tCheck')
datosPotenciales=datosPotenciales.replace('Repair','\n\n\tRepair')
datosPotenciales=datosPotenciales.replace('Line','\n\n\t\tLine')
datosPotenciales=datosPotenciales.replace('\t\t\t','')	

with io.open('/home/jesus/pruebas/archivoACHECKER.text', 'w') as f:
    f.write(datosConocidos +'\n')
    #f.write(datosPosibles +"\n")
    f.write(datosPotenciales+"\n")


document.add_paragraph(datosConocidos)

document.save('dem2o.docx')

driver.close()
